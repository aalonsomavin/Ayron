from datetime import date

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Length, Pt, RGBColor

FOOTER_ATTRIBUTION = "Generado con Ayron"
SPANISH_MONTHS = (
    "ene",
    "feb",
    "mar",
    "abr",
    "may",
    "jun",
    "jul",
    "ago",
    "sep",
    "oct",
    "nov",
    "dic",
)
FONT_FOOTER_PT = 8

FONT_SANS = "Geist"
FONT_SANS_FALLBACK = "Segoe UI"
FONT_BODY_PT = 11
FONT_SUBTITLE_PT = 10
FONT_SECTION_PT = 17
FONT_TITLE_PT = 26
FONT_KICKER_PT = 9
FONT_TABLE_PT = 10

from apps.agent.tools.table_style_tokens import CELL_TONE_COLORS, COLORS

CALLOUT_VARIANTS = {
    "info": {"bg": COLORS["info_bg"], "border": COLORS["info_border"], "label": "Nota"},
    "success": {"bg": COLORS["success_bg"], "border": COLORS["success_border"], "label": "Éxito"},
    "warning": {"bg": COLORS["warning_bg"], "border": COLORS["warning_border"], "label": "Atención"},
    "danger": {"bg": COLORS["danger_bg"], "border": COLORS["danger_border"], "label": "Importante"},
}

ROW_STYLE_FILLS = {
    "default": None,
    "subtotal": COLORS["bg_subtle"],
    "total": COLORS["bg_muted"],
}

PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
PAGE_MARGIN_IN = 0.75
CONTENT_WIDTH_IN = PAGE_WIDTH_IN - (2 * PAGE_MARGIN_IN)
PREVIEW_DPI = 96
SECTION_SPACE_BEFORE = Pt(10)
SECTION_SPACE_AFTER = Pt(6)
BODY_SPACE_AFTER = Pt(4)
BLOCK_SPACE_AFTER = Pt(6)
TABLE_SPACE_BEFORE = Pt(6)
TABLE_SPACE_AFTER = Pt(6)
TABLE_CELL_MARGIN_TWIPS = 50
CALLOUT_CELL_MARGIN_TWIPS = 70


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, *, size_pt: int, bold: bool = False, color: str | None = None):
    run.font.name = FONT_SANS
    run.font.size = Pt(size_pt)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_SANS)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_SANS)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT_SANS)
    if color:
        run.font.color.rgb = rgb(color)


def preview_px(inches: float) -> int:
    return round(inches * PREVIEW_DPI)


def format_document_date(value: date | None = None) -> str:
    current = value or date.today()
    month = SPANISH_MONTHS[current.month - 1]
    return f"{current.day} {month} {current.year}"


def footer_attribution_text(value: date | None = None) -> str:
    return f"{FOOTER_ATTRIBUTION} · {format_document_date(value)}"


def _append_field(paragraph, field_code: str, *, size_pt: int = FONT_FOOTER_PT, color: str = COLORS["text_subtle"]):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    set_run_font(run, size_pt=size_pt, color=color)
    return run


def _clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            element.remove(child)


def configure_document_header(doc, title: str = "", subtitle: str = "", generated_on: date | None = None):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    for paragraph in header.paragraphs:
        _clear_paragraph(paragraph)

    title_text = str(title or "").strip()
    if not title_text:
        return

    table = header.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    table.autofit = False
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    left_cell.width = Inches(4.75)
    right_cell.width = Inches(1.75)

    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = left_paragraph.add_run(title_text)
    set_run_font(title_run, size_pt=FONT_FOOTER_PT, bold=True, color=COLORS["emphasis"])

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_text = str(subtitle or "").strip() or format_document_date(generated_on)
    meta_run = right_paragraph.add_run(meta_text)
    set_run_font(meta_run, size_pt=FONT_FOOTER_PT, color=COLORS["text_subtle"])

    rule_paragraph = header.add_paragraph()
    _clear_paragraph(rule_paragraph)
    rule_paragraph.paragraph_format.space_before = Pt(8)
    rule_paragraph.paragraph_format.space_after = Pt(0)
    p_pr = rule_paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLORS["border"])
    border.append(bottom)
    p_pr.append(border)


def configure_document_footer(doc, generated_on: date | None = None):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    rule_paragraph = footer.paragraphs[0]
    _clear_paragraph(rule_paragraph)
    rule_paragraph.paragraph_format.space_before = Pt(10)
    rule_paragraph.paragraph_format.space_after = Pt(8)
    p_pr = rule_paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), COLORS["border"])
    border.append(top)
    p_pr.append(border)

    table = footer.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    table.autofit = False
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    left_cell.width = Inches(4.75)
    right_cell.width = Inches(1.75)

    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    dot_run = left_paragraph.add_run("● ")
    set_run_font(dot_run, size_pt=7, color=COLORS["accent"])
    attribution_run = left_paragraph.add_run(footer_attribution_text(generated_on))
    set_run_font(attribution_run, size_pt=FONT_FOOTER_PT, color=COLORS["text_subtle"])

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _append_field(right_paragraph, "PAGE")
    de_run = right_paragraph.add_run(" de ")
    set_run_font(de_run, size_pt=FONT_FOOTER_PT, color=COLORS["text_subtle"])
    _append_field(right_paragraph, "NUMPAGES")


def configure_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(PAGE_MARGIN_IN)
    section.bottom_margin = Inches(PAGE_MARGIN_IN)
    section.left_margin = Inches(PAGE_MARGIN_IN)
    section.right_margin = Inches(PAGE_MARGIN_IN)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_SANS
    normal.font.size = Pt(FONT_BODY_PT)
    normal.font.color.rgb = rgb(COLORS["text"])
    normal.paragraph_format.space_after = BODY_SPACE_AFTER
    normal.paragraph_format.line_spacing = 1.35

    for style_name in ("Title", "Heading 1", "Heading 2"):
        _disable_heading_page_breaks(doc.styles[style_name])

    for style_name, size, bold, color in (
        ("Heading 1", FONT_SECTION_PT, True, COLORS["heading"]),
        ("Heading 2", 14, True, COLORS["heading"]),
        ("List Bullet", FONT_BODY_PT, False, COLORS["text"]),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT_SANS
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = rgb(color)


def _set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, *, color: str = COLORS["border"], size: str = "4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    tc_pr.append(borders)


def _set_cell_margins(
    cell,
    *,
    top=TABLE_CELL_MARGIN_TWIPS,
    bottom=TABLE_CELL_MARGIN_TWIPS,
    left=TABLE_CELL_MARGIN_TWIPS + 20,
    right=TABLE_CELL_MARGIN_TWIPS + 20,
):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    tc_pr.append(margins)


def _coerce_spacing(value) -> Length:
    if isinstance(value, Length):
        return value
    return Pt(value)


def _disable_heading_page_breaks(style):
    style.paragraph_format.page_break_before = False
    style.paragraph_format.keep_with_next = False
    style.paragraph_format.keep_together = False


def _reset_heading_paragraph_format(paragraph):
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.keep_together = False


def _style_paragraph(paragraph, *, space_before=0, space_after=BODY_SPACE_AFTER):
    paragraph.paragraph_format.space_before = _coerce_spacing(space_before)
    paragraph.paragraph_format.space_after = _coerce_spacing(space_after)


def add_kicker(doc, text: str):
    paragraph = doc.add_paragraph()
    _style_paragraph(paragraph, space_before=0, space_after=Pt(4))
    run = paragraph.add_run(text.upper())
    set_run_font(run, size_pt=FONT_KICKER_PT, bold=True, color=COLORS["accent"])
    paragraph.paragraph_format.space_after = Pt(4)


def add_document_header(doc, title: str, subtitle: str = ""):
    add_title(doc, title)
    if subtitle.strip():
        add_subtitle(doc, subtitle.strip())
    add_separator(doc)


def add_title(doc, text: str):
    paragraph = doc.add_heading(text, level=0)
    _reset_heading_paragraph_format(paragraph)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size_pt=FONT_TITLE_PT, bold=True, color=COLORS["heading"])


def add_subtitle(doc, text: str):
    paragraph = doc.add_paragraph()
    _style_paragraph(paragraph, space_after=Pt(10))
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=FONT_SUBTITLE_PT, color=COLORS["text_muted"])


def add_section_heading(doc, text: str):
    paragraph = doc.add_heading(text, level=1)
    _reset_heading_paragraph_format(paragraph)
    paragraph.paragraph_format.space_before = SECTION_SPACE_BEFORE
    paragraph.paragraph_format.space_after = SECTION_SPACE_AFTER
    for run in paragraph.runs:
        set_run_font(run, size_pt=FONT_SECTION_PT, bold=True, color=COLORS["heading"])


def add_body_paragraph(doc, text: str):
    paragraph = doc.add_paragraph()
    _style_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=FONT_BODY_PT, color=COLORS["text"])


def add_bullet_item(doc, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    _style_paragraph(paragraph, space_after=Pt(4))
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=FONT_BODY_PT, color=COLORS["text"])


def add_separator(doc):
    paragraph = doc.add_paragraph()
    _style_paragraph(paragraph, space_before=Pt(6), space_after=Pt(6))
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLORS["border"])
    border.append(bottom)
    p_pr.append(border)


def _set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    tr_pr.append(header)


def _paragraph_alignment(align: str):
    if align == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if align == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT


def _render_table_cell(
    cell,
    *,
    cell_data: dict,
    row_style: str,
    row_idx: int,
):
    value = cell_data.get("value", "")
    tone = cell_data.get("tone", "default")
    align = cell_data.get("align", "left")
    bold = bool(cell_data.get("bold", False))
    row_fill = ROW_STYLE_FILLS.get(row_style)
    if row_fill:
        fill = row_fill
    else:
        fill = COLORS["bg_subtle"] if row_idx % 2 == 0 else "FFFFFF"
    text_color = CELL_TONE_COLORS.get(tone, COLORS["text"])
    border_color = COLORS["border"] if row_style == "total" else COLORS["border_subtle"]
    border_size = "8" if row_style == "total" else "4"

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = _paragraph_alignment(align)
    run = paragraph.add_run(str(value))
    set_run_font(run, size_pt=FONT_TABLE_PT, bold=bold, color=text_color)
    _set_cell_shading(cell, fill)
    _set_cell_border(cell, color=border_color, size=border_size)
    _set_cell_margins(cell)


def _table_body_rows(rows: list) -> list[tuple[str, list]]:
    body_rows = []
    for row in rows:
        if isinstance(row, dict):
            body_rows.append((row.get("style", "default"), row.get("cells", [])))
        elif isinstance(row, list):
            body_rows.append(("default", row))
    return body_rows


def _coerce_table_cell_data(cell, *, default_bold: bool = False) -> dict:
    from apps.agent.tools.report_content import normalize_docx_cell

    return normalize_docx_cell(cell, default_bold=default_bold)


def add_styled_table(doc, table: dict):
    headers = table["headers"]
    rows = _table_body_rows(table["rows"])
    caption = str(table.get("caption", "")).strip()

    if caption:
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.paragraph_format.space_before = Pt(0)
        caption_paragraph.paragraph_format.space_after = Pt(2)
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(caption_run, size_pt=FONT_TABLE_PT, bold=True, color=COLORS["emphasis"])

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = TABLE_SPACE_BEFORE

    table_el = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table_el.autofit = True
    table_el.allow_autofit = True

    for col_idx, header in enumerate(headers):
        cell = table_el.rows[0].cells[col_idx]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        set_run_font(run, size_pt=FONT_TABLE_PT, bold=True, color=COLORS["emphasis"])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(cell, COLORS["bg_muted"])
        _set_cell_border(cell, color=COLORS["border"], size="4")
        _set_cell_margins(cell)

    _set_row_repeat_header(table_el.rows[0])

    for row_idx, (row_style, cells) in enumerate(rows):
        default_bold = row_style in ("total", "subtotal")
        for col_idx, _header in enumerate(headers):
            cell = table_el.rows[row_idx + 1].cells[col_idx]
            raw_cell = cells[col_idx] if col_idx < len(cells) else ""
            cell_data = _coerce_table_cell_data(raw_cell, default_bold=default_bold)
            _render_table_cell(
                cell,
                cell_data=cell_data,
                row_style=row_style,
                row_idx=row_idx,
            )

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = BLOCK_SPACE_AFTER


def add_callout(doc, *, variant: str, title: str, text: str):
    theme = CALLOUT_VARIANTS.get(variant, CALLOUT_VARIANTS["info"])
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    _set_cell_shading(cell, theme["bg"])
    pad = CALLOUT_CELL_MARGIN_TWIPS
    _set_cell_margins(cell, top=pad, bottom=pad, left=pad + 20, right=pad + 20)

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "right", "bottom"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), COLORS["border_subtle"])
        borders.append(element)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), theme["border"])
    borders.append(left)
    tc_pr.append(borders)

    label = title.strip() or theme["label"]
    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run(label)
    set_run_font(title_run, size_pt=FONT_TABLE_PT, bold=True, color=COLORS["emphasis"])
    title_paragraph.paragraph_format.space_after = Pt(4)

    body_paragraph = cell.add_paragraph()
    body_run = body_paragraph.add_run(text)
    set_run_font(body_run, size_pt=FONT_BODY_PT, color=COLORS["text_muted"])
    body_paragraph.paragraph_format.space_after = Pt(0)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = BLOCK_SPACE_AFTER
